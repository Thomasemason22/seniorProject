import os
import json
import sqlite3
from collections import defaultdict

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


BASE_DIR = os.path.abspath(os.path.dirname(__file__))
ROOT = os.path.abspath(os.path.join(BASE_DIR, "..", ".."))
DB_PATH = os.path.join(ROOT, "instance", "database.db")
ASSET_DIR = os.path.join(BASE_DIR, "assets")
OUT_PATH = os.path.join(BASE_DIR, "OutboundOps_Senior_Project_Report.docx")
METRICS_PATH = os.path.join(BASE_DIR, "metrics.json")

INK = RGBColor(45, 24, 15)
BROWN = RGBColor(53, 28, 21)
GOLD = RGBColor(255, 181, 0)
BLUE = RGBColor(59, 95, 138)
GREEN = RGBColor(23, 114, 69)
MUTED = RGBColor(114, 90, 73)
LINE = "E4D5C2"
FILL = "FBF7EF"


def moneyless_number(value, decimals=0):
    if decimals:
        return f"{value:,.{decimals}f}"
    return f"{round(value):,}"


def pct(value, decimals=1):
    return f"{value * 100:.{decimals}f}%"


def load_rows():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    rows = [dict(row) for row in conn.execute("select * from warehouse_operation")]
    conn.close()
    return rows


def metrics(rows):
    total_volume = sum(row["gross_volume"] or 0 for row in rows)
    outbound_volume = sum((row["gross_volume"] or 0) for row in rows if row["area_group"] == "Outbounds")
    scanned = sum(row["scanned_volume"] or 0 for row in rows)
    paid = sum(row["paid_day"] or 0 for row in rows)
    planned = sum(row["planned_hours"] or 0 for row in rows)
    overtime = sum(row["overtime_hours"] or 0 for row in rows)
    sort_totals = defaultdict(lambda: {"volume": 0, "paid": 0, "ot": 0})
    groups = defaultdict(lambda: {"records": 0, "volume": 0, "paid": 0, "ot": 0, "pph": []})
    belts = defaultdict(lambda: {"records": 0, "volume": 0, "scan": 0, "paid": 0})

    for row in rows:
        key = (row["date"], row["shift"])
        sort_totals[key]["volume"] += row["gross_volume"] or 0
        sort_totals[key]["paid"] += row["paid_day"] or 0
        sort_totals[key]["ot"] += row["overtime_hours"] or 0
        group = groups[row["area_group"]]
        group["records"] += 1
        group["volume"] += row["gross_volume"] or 0
        group["paid"] += row["paid_day"] or 0
        group["ot"] += row["overtime_hours"] or 0
        group["pph"].append(row["actual_pph"] or 0)
        belt = belts[row["belt"]]
        belt["records"] += 1
        belt["volume"] += row["gross_volume"] or 0
        belt["scan"] += row["scanned_volume"] or 0
        belt["paid"] += row["paid_day"] or 0

    sort_values = [item["volume"] for item in sort_totals.values()]
    group_rows = []
    for name, item in groups.items():
        group_rows.append(
            {
                "name": name,
                "records": item["records"],
                "volume": item["volume"],
                "share": item["volume"] / total_volume if total_volume else 0,
                "avg_pph": sum(item["pph"]) / len(item["pph"]) if item["pph"] else 0,
                "overtime": item["ot"],
            }
        )
    group_rows.sort(key=lambda item: item["volume"], reverse=True)

    belt_rows = []
    for name, item in belts.items():
        belt_rows.append(
            {
                "name": name,
                "volume": item["volume"],
                "scan_rate": item["scan"] / item["volume"] if item["volume"] and item["scan"] else 0,
                "packages_per_paid_hour": item["volume"] / item["paid"] if item["paid"] else 0,
            }
        )
    belt_rows.sort(key=lambda item: item["volume"], reverse=True)

    return {
        "records": len(rows),
        "date_from": min(row["date"] for row in rows),
        "date_to": max(row["date"] for row in rows),
        "total_volume": total_volume,
        "outbound_volume": outbound_volume,
        "scanned": scanned,
        "scan_rate": scanned / outbound_volume if outbound_volume else 0,
        "paid": paid,
        "planned": planned,
        "overtime": overtime,
        "avg_pph": sum(row["actual_pph"] or 0 for row in rows) / len(rows),
        "avg_staffing": sum(row["staffing_level"] or 0 for row in rows) / len(rows),
        "sort_count": len(sort_values),
        "avg_sort_volume": sum(sort_values) / len(sort_values),
        "peak_sort_volume": max(sort_values),
        "groups": group_rows,
        "belts": belt_rows,
    }


def try_font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for path in candidates:
        if path and os.path.exists(path):
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def draw_group_chart(data, path):
    width, height = 1500, 820
    im = Image.new("RGB", (width, height), "#fbf7ef")
    draw = ImageDraw.Draw(im)
    title = try_font(44, True)
    label = try_font(28, True)
    small = try_font(24)
    draw.text((70, 48), "Volume by Operating Area", fill="#2d180f", font=title)
    draw.text((70, 105), "Seeded 30-day warehouse operations dataset", fill="#725a49", font=small)
    chart_x, chart_y = 360, 170
    bar_h, gap = 54, 24
    max_value = max(item["volume"] for item in data)
    colors = ["#351c15", "#ffb500", "#3b5f8a", "#177245", "#b42318", "#7a5a00", "#725a49", "#a36a2c"]
    for idx, item in enumerate(data):
        y = chart_y + idx * (bar_h + gap)
        draw.text((70, y + 10), item["name"], fill="#2d180f", font=label)
        draw.rounded_rectangle((chart_x, y, width - 150, y + bar_h), radius=12, fill="#ead8bd")
        bar_w = int((width - 150 - chart_x) * item["volume"] / max_value)
        draw.rounded_rectangle((chart_x, y, chart_x + bar_w, y + bar_h), radius=12, fill=colors[idx % len(colors)])
        draw.text((chart_x + bar_w + 18, y + 10), f"{item['share']*100:.1f}% | {moneyless_number(item['volume'])}", fill="#2d180f", font=small)
    im.save(path, quality=95)


def draw_architecture(path):
    width, height = 1500, 740
    im = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(im)
    title = try_font(44, True)
    h = try_font(31, True)
    body = try_font(24)
    draw.text((70, 48), "System Architecture", fill="#2d180f", font=title)
    boxes = [
        ("Seed Data", "30 days x 3 shifts\nUL, PD, SRT, Metro,\nand support areas", "#fbf7ef"),
        ("Flask API", "CRUD endpoints\nKPI endpoint\nSQLite persistence", "#fff7d6"),
        ("Analytics Layer", "Pandas KPI rollups\nscan rate, PPH,\npaid day, risk inputs", "#eef5ff"),
        ("React Dashboard", "filters, charts,\ntables, export,\nprintable report view", "#effaf3"),
    ]
    x, y, box_w, box_h = 70, 170, 300, 310
    for idx, (name, detail, fill) in enumerate(boxes):
        bx = x + idx * 355
        draw.rounded_rectangle((bx, y, bx + box_w, y + box_h), radius=18, fill=fill, outline="#e4d5c2", width=4)
        draw.text((bx + 28, y + 34), name, fill="#351c15", font=h)
        draw.multiline_text((bx + 28, y + 102), detail, fill="#725a49", font=body, spacing=11)
        if idx < len(boxes) - 1:
            ax = bx + box_w + 32
            ay = y + box_h // 2
            draw.line((ax, ay, ax + 42, ay), fill="#ffb500", width=10)
            draw.polygon([(ax + 42, ay), (ax + 18, ay - 18), (ax + 18, ay + 18)], fill="#ffb500")
    draw.text((70, 560), "Design goal: convert operational logs into a supervisor-ready control center for staffing, flow, and outbound scan decisions.", fill="#2d180f", font=body)
    im.save(path, quality=95)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="DADCE0"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_shading(header_cells[idx], FILL)
        set_cell_border(header_cells[idx], LINE)
        header_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in header_cells[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = BROWN
                run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            set_cell_border(cells[idx], LINE)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = INK
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.color.rgb = BROWN if level == 1 else BLUE
        run.font.name = "Calibri"
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        run.font.color.rgb = INK


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        run.font.color.rgb = INK


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = INK
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    for style_name, size, color in [
        ("Title", 24, BROWN),
        ("Heading 1", 16, BROWN),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, BLUE),
    ]:
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = color
    header = section.header.paragraphs[0]
    header.text = "OutboundOps Dashboard | Senior Project"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.text = "Generated from project source files and SQLite seed dataset"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def build_report():
    rows = load_rows()
    m = metrics(rows)
    with open(METRICS_PATH, "w", encoding="utf-8") as handle:
        json.dump(m, handle, indent=2)
    os.makedirs(ASSET_DIR, exist_ok=True)
    group_chart = os.path.join(ASSET_DIR, "volume_by_area.png")
    arch_chart = os.path.join(ASSET_DIR, "architecture.png")
    draw_group_chart(m["groups"], group_chart)
    draw_architecture(arch_chart)

    doc = Document()
    configure_doc(doc)

    doc.add_paragraph("SENIOR PROJECT REPORT").runs[0].font.color.rgb = MUTED
    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    run = title.add_run("OutboundOps Dashboard")
    run.bold = True
    run.font.color.rgb = BROWN
    subtitle = doc.add_paragraph("A Flask and React analytics dashboard for warehouse outbound operations")
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = MUTED
    meta = doc.add_paragraph(f"Prepared: May 10, 2026 | Dataset window: {m['date_from']} to {m['date_to']}")
    meta.runs[0].font.size = Pt(10.5)
    meta.runs[0].font.color.rgb = MUTED
    doc.add_paragraph()

    add_table(
        doc,
        ["Metric", "Value", "Why it matters"],
        [
            ["Records analyzed", moneyless_number(m["records"]), "Confirms broad enough test coverage for dashboard behavior."],
            ["Total gross volume", moneyless_number(m["total_volume"]), "Represents building-level package flow over the seeded period."],
            ["Outbound scan rate", pct(m["scan_rate"]), "Shows how scanned PD activity compares with outbound gross volume."],
            ["Avg shift volume", moneyless_number(m["avg_sort_volume"]), "Provides the baseline for staffing and capacity decisions."],
            ["Peak shift volume", moneyless_number(m["peak_sort_volume"]), "Identifies high-pressure sorts that need planning attention."],
            ["Total paid day", moneyless_number(m["paid"], 1), "Summarizes labor hours consumed across operations."],
        ],
        widths=[1.7, 1.35, 3.25],
    )

    add_heading(doc, "Executive Summary")
    doc.add_paragraph(
        "OutboundOps Dashboard is a full-stack senior project that turns warehouse operations records into a supervisor-oriented control center. "
        "The application combines a Flask API, SQLite persistence, a React dashboard, KPI calculations, chart panels, filtering tools, risk scoring, "
        "recommendations, CSV export, and a printable daily-report workflow."
    )
    doc.add_paragraph(
        f"The seeded evaluation dataset contains {moneyless_number(m['records'])} operational records across {moneyless_number(m['sort_count'])} shifts. "
        f"It models {moneyless_number(m['total_volume'])} total packages, with Outbounds carrying {pct(m['outbound_volume'] / m['total_volume'])} of volume "
        f"and reaching an {pct(m['scan_rate'])} scan rate. These figures make the project useful for demonstrating data mining ideas: aggregation, filtering, "
        "trend analysis, anomaly detection, operational risk signals, and decision support."
    )

    add_heading(doc, "Problem and Motivation")
    doc.add_paragraph(
        "Warehouse supervisors often need to connect package volume, scanning progress, staffing, paid day, overtime, and area performance quickly enough "
        "to adjust the current or next shift. Raw operational tables are useful for recordkeeping, but they make it difficult to see where labor is overloaded, "
        "which PD belts are under-scanning, or whether peak volume is beginning to exceed staffing coverage."
    )
    add_bullets(
        doc,
        [
            "Centralize operational records for unload, outbounds, airsort, sort aisle, small sort, irregulars, indirect, and metro areas.",
            "Summarize performance through KPIs such as gross volume, scanned volume, paid day, overtime, planned hours, throughput, and PPH.",
            "Expose decision-ready views for shift comparison, belt detail, forecast, anomaly review, risk scoring, and recommendations.",
            "Support practical supervisor workflows such as filtering, search, CSV export, bulk record entry, and printing a daily report.",
        ],
    )

    add_heading(doc, "Data Model and Seeded Dataset")
    doc.add_paragraph(
        "The backend stores each operational observation as a WarehouseOperation record. The table includes date, shift, area group, belt, gross volume, "
        "scanned volume, staffing level, hours, paid day, throughput, actual PPH, planned PPH, planned hours, overtime, and notes. The seed script creates "
        "a controlled synthetic dataset for 30 days and three shifts per day."
    )
    add_table(
        doc,
        ["Area group", "Records", "Gross volume", "Volume share", "Avg PPH", "Overtime"],
        [
            [
                item["name"],
                moneyless_number(item["records"]),
                moneyless_number(item["volume"]),
                pct(item["share"]),
                moneyless_number(item["avg_pph"], 1),
                moneyless_number(item["overtime"], 1),
            ]
            for item in m["groups"]
        ],
        widths=[1.25, 0.8, 1.25, 1.0, 0.9, 0.95],
    )
    doc.add_picture(group_chart, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "System Design")
    doc.add_paragraph(
        "The project uses a conventional but effective full-stack architecture. Flask owns the data API and KPI endpoint, SQLAlchemy maps operations to SQLite, "
        "and React owns the interactive dashboard. The frontend requests data from the API, derives filtered summaries, and passes filtered records into reusable "
        "visual components."
    )
    doc.add_picture(arch_chart, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table(
        doc,
        ["Layer", "Implementation", "Responsibilities"],
        [
            ["Backend API", "Flask, Flask-CORS, SQLAlchemy", "Expose operations endpoints, create single/bulk records, calculate KPIs, and serve JSON."],
            ["Storage", "SQLite database", "Persist seeded and user-created warehouse operation records."],
            ["Analytics", "Pandas in analytics.py plus frontend reducers", "Aggregate volume, scanned volume, throughput, staffing, paid day, overtime, and planned hours."],
            ["Frontend", "React, Chart.js, component panels", "Render dashboard filters, KPI cards, charts, risk panels, recommendations, export, and report views."],
        ],
        widths=[1.1, 1.65, 3.45],
    )

    add_heading(doc, "Analytics and Data Mining Methods")
    doc.add_paragraph(
        "The project is intentionally focused on applied operational analytics rather than a black-box model. The data mining value comes from turning granular "
        "records into interpretable features and decision signals."
    )
    add_numbered(
        doc,
        [
            "Aggregation: group records by shift, area group, belt, and date range to calculate total volume, scanned volume, paid day, overtime, staffing, and PPH.",
            "Filtering and slicing: allow users to isolate one shift, shift type, area group, belt, date range, or search term, then recompute all visible metrics.",
            "Performance benchmarking: compare planned hours against paid day and compare area throughput against average operating performance.",
            "Risk scoring: combine high-overtime exposure with throughput pressure to label workload risk as stable, moderate, or elevated.",
            "Recommendation logic: generate next-action suggestions such as flex coverage, staffing-floor review, and area-ranking standup focus.",
            "Forecast and anomaly views: use historical records to highlight unusual volume or labor patterns before they become staffing problems.",
        ],
    )

    add_heading(doc, "Dashboard Features")
    add_bullets(
        doc,
        [
            "KPI cards summarize total volume, average sort volume, peak sort volume, average PPH, staffing, paid day, planned hours, and overtime.",
            "Chart panels cover volume trends, staffing, area performance, flow comparison, shift mix, PPH distribution, capacity, and outbound load rate.",
            "Operational panels expose PD belt heatmaps, belt detail, area staff detail, alerts, anomaly flags, area rankings, and recommended actions.",
            "Shift tools support single-record creation, bulk record creation, CSV export, and printable report output for non-viewer roles.",
            "Filters apply globally, making the dashboard responsive to supervisor questions rather than fixed static reports.",
        ],
    )

    add_heading(doc, "Evaluation")
    doc.add_paragraph(
        "Evaluation was performed by inspecting the source code, calculating metrics directly from the SQLite seed database, and confirming that the dataset can "
        "support the major dashboard views. The metrics show that the seeded data reflects realistic operational variety: Outbounds is the dominant workstream, "
        "Unload and Sort Aisle create meaningful secondary load, and smaller areas remain visible without overwhelming the dashboard."
    )
    add_table(
        doc,
        ["Evaluation question", "Evidence from current project", "Result"],
        [
            ["Does the app cover the main warehouse areas?", "Seed data includes UL 1-6, PD 1-18, SRT 1-6, Metro 1-4, Airsort, Small Sort, Irregulars, and Indirect.", "Met"],
            ["Can supervisors isolate operational slices?", "Dashboard filters support specific shift, shift type, group, area, date range, and search.", "Met"],
            ["Are labor and flow tied together?", "Records store volume, staffing, hours, paid day, overtime, throughput, planned hours, and PPH.", "Met"],
            ["Does the system go beyond descriptive totals?", "Risk, recommendations, forecast, anomaly, ranking, and capacity panels add decision support.", "Met"],
            ["Is the dataset large enough for demo analysis?", f"{moneyless_number(m['records'])} records across {moneyless_number(m['sort_count'])} shifts.", "Met"],
        ],
        widths=[1.65, 3.45, 0.75],
    )

    add_heading(doc, "Limitations")
    add_bullets(
        doc,
        [
            "The dataset is synthetic; it is useful for demonstration but not a substitute for real warehouse telemetry.",
            "Forecasting and anomaly logic should be validated against actual historical operations before production use.",
            "Authentication and role permissions are represented in the UI but would need a secure backend implementation for deployment.",
            "The SQLite database is appropriate for a senior-project prototype; a production system would likely use PostgreSQL or an enterprise data warehouse.",
            "The current project calculates many summary statistics on the client, which is acceptable for this dataset size but should move server-side for larger deployments.",
        ],
    )

    add_heading(doc, "Future Work")
    add_bullets(
        doc,
        [
            "Connect to real operational feeds or scheduled CSV imports.",
            "Add authenticated supervisor, admin, and viewer accounts backed by the API.",
            "Train predictive models for volume, staffing need, and scan-rate risk by shift and area.",
            "Add alert thresholds that can be configured by building or operating plan.",
            "Create a nightly report export that emails supervisors a PDF or spreadsheet summary.",
        ],
    )

    add_heading(doc, "Conclusion")
    doc.add_paragraph(
        "OutboundOps Dashboard demonstrates how data mining techniques can be applied to a practical warehouse operations problem. The project does not simply "
        "store data; it transforms records into KPIs, charts, rankings, risk indicators, and action recommendations. As a senior project, it shows full-stack "
        "implementation ability, data modeling, analytics thinking, and a clear understanding of how software can support operational decision-making."
    )

    doc.add_page_break()
    add_heading(doc, "Appendix: Key Source Files")
    add_table(
        doc,
        ["File", "Purpose"],
        [
            ["app.py", "Defines the Flask API, database configuration, operations endpoints, bulk insert, and KPI endpoint."],
            ["models.py", "Defines the WarehouseOperation SQLAlchemy model and JSON serialization."],
            ["analytics.py", "Calculates backend KPI summaries from stored operation records."],
            ["seed_data.py", "Creates the synthetic warehouse operations dataset used for the dashboard demo."],
            ["frontend/src/Dashboard.js", "Coordinates dashboard state, API calls, filters, summaries, role controls, export, and report printing."],
            ["frontend/src/components/*", "Contains reusable chart and analysis panels for operations, staffing, risk, recommendations, and belt detail."],
        ],
        widths=[2.0, 4.2],
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_report())
